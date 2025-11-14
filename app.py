# app.py — Polished AI Resume Builder
from dotenv import load_dotenv
import os
load_dotenv()  # must run before os.getenv()
import streamlit as st

OPENAI_KEY = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
GEMINI_KEY = st.secrets.get("GEMINI_API_KEY") or os.getenv("GEMINI_API_KEY")

import io
import random
import requests
import subprocess
import tempfile
from typing import Tuple, List, Dict
from PyPDF2 import PdfReader
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# AI libs
import google.generativeai as genai
from openai import OpenAI

# -------------------------
# CONFIG
# -------------------------
st.set_page_config(page_title="AI Resume Agent", page_icon="🤖", layout="wide")
st.title("🤖 AI Resume Builder & ATS Optimization Agent")
st.markdown("Upload or paste your resume, enhance with AI (OpenAI / Gemini) and download ATS-optimized DOCX/PDF.")

if GEMINI_KEY:
    genai.configure(api_key=GEMINI_KEY)

st.sidebar.markdown("### API Keys")
st.sidebar.write("OpenAI:", "✅" if OPENAI_KEY else "❌")
st.sidebar.write("Gemini:", "✅" if GEMINI_KEY else "❌")

# -------------------------
# Caching AI calls
# -------------------------
@st.cache_data(show_spinner=False)
def cached_ai_score(resume_text: str, use_gemini_flag: bool) -> float:
    """Compute AI-based ATS score (cached)."""
    return ats_score_ai(resume_text, use_gemini_flag)

@st.cache_data(show_spinner=False)
def cached_ai_enhance(resume_text: str, prompt: str, use_gemini_flag: bool) -> str:
    """Enhance resume text with AI (cached)."""
    return call_ai(resume_text, prompt, use_gemini_flag)

# -------------------------
# Utilities & Helpers
# -------------------------
KEYWORDS = [
    "Python", "Machine Learning", "AI", "Data", "Developer", "Analysis", "Engineer",
    "TensorFlow", "Pandas", "NumPy", "Scikit-learn", "Streamlit", "SQL"
]

WEIGHTS = {
    "python": 3, "machine learning": 4, "ai": 3, "data": 2,
    "developer": 2, "analysis": 2, "engineer": 3,
    "tensorflow": 2, "pandas": 2, "numpy": 2, "scikit-learn": 2, "streamlit": 2, "sql": 2
}

def safe_extract_text_from_pdf(uploaded_file) -> str:
    text = ""
    try:
        reader = PdfReader(uploaded_file)
        for p in reader.pages:
            text += (p.extract_text() or "") + "\n"
    except Exception:
        return ""
    return text

def get_ats_score_local(text: str) -> float:
    t = (text or "").lower()
    score = 0
    for kw, w in WEIGHTS.items():
        hits = min(t.count(kw.lower()), 3)
        score += hits * w
    max_score = sum(WEIGHTS.values()) * 3
    return round((score / max_score) * 100, 2) if max_score else 0.0

def get_ats_score_api(text: str) -> float:
    try:
        resp = requests.post("https://mock-ats-api.fly.dev/score", json={"resume_text": text}, timeout=4)
        if resp.status_code == 200:
            return float(resp.json().get("score", 0))
    except Exception:
        pass
    return get_ats_score_local(text)

def ats_score_ai(resume_text: str, use_gemini_flag: bool) -> float:
    prompt = f"Evaluate the following resume for ATS keywords and give a score out of 100:\n\n{resume_text}"
    try:
        resp_text = call_ai("", prompt, use_gemini_flag)
        return float(resp_text.strip().split()[0])
    except Exception:
        return get_ats_score_local(resume_text)

def re_inject_keywords(orig: str, enhanced: str, keywords=KEYWORDS) -> Tuple[str, List[str]]:
    missing = [k for k in keywords if k.lower() in (orig or "").lower() and k.lower() not in (enhanced or "").lower()]
    if missing:
        enhanced = (enhanced or "").strip() + "\n\nTechnical Keywords: " + ", ".join(missing)
    return enhanced, missing

def generate_docx_from_text(text: str) -> io.BytesIO:
    doc = Document()
    doc.add_heading("Resume (AI-Enhanced)", level=0)
    for line in (text or "").split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def generate_pdf_from_text(text: str) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elems = []
    for para in (text or "").split("\n"):
        if para.strip():
            elems.append(Paragraph(para.strip().replace("&", "&amp;"), styles["Normal"]))
            elems.append(Spacer(1, 6))
    doc.build(elems)
    buffer.seek(0)
    return buffer

# -------------------------
# AI Enhancement wrappers
# -------------------------
def enhance_with_openai(text: str, prompt: str) -> str:
    if not OPENAI_KEY:
        raise RuntimeError("OpenAI API key not configured.")
    client = OpenAI(api_key=OPENAI_KEY)
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt + "\n\n" + (text or "")}],
        temperature=0.2,
        max_tokens=1200
    )
    return resp.choices[0].message.content

def enhance_with_gemini(text: str, prompt: str) -> str:
    if not GEMINI_KEY:
        raise RuntimeError("Gemini API key not configured.")
    for m in ("models/gemini-2.5-flash", "models/gemini-2.5-pro", "models/gemini-flash-latest"):
        try:
            model = genai.GenerativeModel(m)
            resp = model.generate_content(prompt + "\n\n" + (text or ""))
            return resp.text
        except Exception:
            continue
    raise RuntimeError("No working Gemini model available.")

def call_ai(text: str, prompt: str, use_gemini_flag: bool) -> str:
    if use_gemini_flag:
        return enhance_with_gemini(text, prompt)
    else:
        return enhance_with_openai(text, prompt)

# -------------------------
# LaTeX helpers
# -------------------------
TEMPLATE_MAP = {
    "ModernCV (LaTeX)": "moderncv.tex",
    "AutoCV (LaTeX)": "autocv.tex"
}

def extract_sections(enhanced_text: str) -> Dict[str, str]:
    out = {k: "" for k in ("NAME","EMAIL","PHONE","SUMMARY","SKILLS","EXPERIENCE","PROJECTS","EDUCATION","LINKEDIN","GITHUB")}
    txt = (enhanced_text or "").strip()
    if not txt:
        return out
    lines = [l.strip() for l in txt.splitlines() if l.strip()]
    out["NAME"] = lines[0] if lines else ""
    for l in lines[:8]:
        if "@" in l and not out["EMAIL"]: out["EMAIL"] = l
        if any(ch.isdigit() for ch in l) and len([c for c in l if c.isdigit()]) >= 6 and not out["PHONE"]: out["PHONE"] = l
        if "linkedin" in l.lower() and not out["LINKEDIN"]: out["LINKEDIN"] = l
        if "github" in l.lower() and not out["GITHUB"]: out["GITHUB"] = l
    lower = txt.lower()
    def grab(section_name):
        s = section_name.lower()
        if s in lower:
            start = lower.index(s)
            end = lower.find("\n\n", start)
            return txt[start:end].strip() if end != -1 else txt[start:].strip()
        return ""
    out["SUMMARY"] = grab("summary").split(":",1)[-1].strip()
    out["SKILLS"] = grab("skills").split(":",1)[-1].strip()
    out["EXPERIENCE"] = grab("experience").split(":",1)[-1].strip()
    out["PROJECTS"] = grab("projects").split(":",1)[-1].strip()
    out["EDUCATION"] = grab("education").split(":",1)[-1].strip()
    return out

def fill_latex_template(template_path: str, data: Dict[str,str]) -> str:
    with open(template_path, "r", encoding="utf-8") as f:
        tex = f.read()
    for k, v in data.items():
        tex = tex.replace(f"<{k}>", v or "")
    return tex

def build_latex_resume(template_name: str, data: Dict[str,str]) -> str:
    template_path = os.path.join("templates", template_name)
    if not os.path.exists(template_path):
        st.error(f"Template file not found: {template_path}")
        return None
    tex_content = fill_latex_template(template_path, data)
    tmp_dir = tempfile.mkdtemp()
    tex_file = os.path.join(tmp_dir, "resume.tex")
    with open(tex_file, "w", encoding="utf-8") as f: f.write(tex_content)
    try:
        subprocess.run(["pdflatex","-interaction=nonstopmode","resume.tex"], cwd=tmp_dir, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
        subprocess.run(["pdflatex","-interaction=nonstopmode","resume.tex"], cwd=tmp_dir, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
    except subprocess.CalledProcessError:
        st.error("LaTeX compilation failed. Falling back to simple PDF.")
        return None
    pdf_path = os.path.join(tmp_dir, "resume.pdf")
    return pdf_path if os.path.exists(pdf_path) else None

# -------------------------
# Session state init (safe)
# -------------------------
# Strings
for key in ["resume_text", "enhanced_text", "resume_editor"]:
    if key not in st.session_state:
        st.session_state[key] = ""

# Lists
if "score_history" not in st.session_state or not isinstance(st.session_state["score_history"], list):
    st.session_state.score_history = []

if "ai_scores" not in st.session_state or not isinstance(st.session_state["ai_scores"], dict):
    st.session_state.ai_scores = {}

# Optional: missing keywords
if "enhance_keywords_missing" not in st.session_state:
    st.session_state.enhance_keywords_missing = []


# -------------------------
# Choose AI engine
# -------------------------
use_gemini = st.checkbox("Use Gemini (Google) instead of OpenAI", value=False)

# -------------------------
# Step 1 — Upload / Manual
# -------------------------
st.header("Step 1 — Upload or Enter Resume")
col1, col2 = st.columns([2,1])
with col1:
    method = st.radio("Input method:", ["Upload PDF/DOCX", "Paste / Manual Form"])
    if method == "Upload PDF/DOCX":
        uploaded = st.file_uploader("Upload resume (PDF or DOCX)", type=["pdf","docx"])
        if uploaded:
            if uploaded.name.lower().endswith(".pdf"):
                extracted = safe_extract_text_from_pdf(uploaded)
            else:
                try:
                    doc = Document(uploaded)
                    extracted = "\n".join(p.text for p in doc.paragraphs)
                except Exception:
                    extracted = ""
            if extracted.strip():
                st.session_state.resume_text = extracted
    else:
        # manual form
        name = st.text_input("Full name", "")
        email = st.text_input("Email", "")
        phone = st.text_input("Phone", "")
        linkedin = st.text_input("LinkedIn / profile", "")
        github = st.text_input("GitHub / profile", "")
        edu = st.text_area("Education (brief)", height=80)
        skills = st.text_area("Skills (comma or bullet list)", height=80)
        work = st.text_area("Work experience (paste bullets)", height=120)
        projects = st.text_area("Projects / achievements", height=100)
        summary = st.text_area("Summary / Objective (optional)", height=80)
        if st.button("✅ Merge manual entries into Resume Text"):
            parts = []
            if name: parts.append(name)
            if email: parts.append(email)
            if phone: parts.append(phone)
            if linkedin: parts.append(linkedin)
            if github: parts.append(github)
            if summary: parts.append("\nSummary:\n" + summary)
            if edu: parts.append("\nEducation:\n" + edu)
            if skills: parts.append("\nSkills:\n" + skills)
            if work: parts.append("\nExperience:\n" + work)
            if projects: parts.append("\nProjects:\n" + projects)
            st.session_state.resume_text = "\n".join(parts)
            st.success("Manual inputs copied into resume text editor below.")

with col2:
    st.subheader("Original ATS Score (live)")
    live_score = get_ats_score_local(st.session_state.resume_text)
    st.metric("Original ATS Score", f"{live_score} / 100")
    if st.session_state.score_history:
        st.line_chart([r["final"] for r in st.session_state.score_history])

if not st.session_state.resume_text.strip():
    st.info("Upload a resume or paste / create one using the manual form to continue.")
    st.stop()

# show extracted text
st.subheader("Resume Text Editor")
st.text_area("Resume Text (you can edit)", value=st.session_state.resume_text, height=260, key="resume_editor")
st.session_state.resume_text = st.session_state.get("resume_editor") or ""

# -------------------------
# Step 2 — Compute AI ATS Score (optimized)
# -------------------------
st.header("Step 2 — Compute AI-based ATS Score")

if "ai_scores" not in st.session_state:
    st.session_state.ai_scores = {}

resume_key = f"score_{st.session_state.resume_text}_{use_gemini}"

if st.button("Compute AI ATS Score"):
    with st.spinner("Calculating AI score..."):
        if resume_key not in st.session_state.ai_scores:
            try:
                st.session_state.ai_scores[resume_key] = cached_ai_score(st.session_state.resume_text, use_gemini)
            except Exception as e:
                st.error(f"AI score failed: {e}")
                st.session_state.ai_scores[resume_key] = get_ats_score_local(st.session_state.resume_text)

local_score = get_ats_score_local(st.session_state.resume_text)
ai_score = st.session_state.ai_scores.get(resume_key, local_score)

st.metric("Local ATS Score", f"{local_score} / 100")
st.metric("AI-based ATS Score", f"{ai_score} / 100", delta=(ai_score - local_score))

# -------------------------
# Step 3 — Enhance Resume (optimized)
# -------------------------
enhance_prompt = """You are an expert technical resume writer and ATS specialist.
Enhance the resume text to improve clarity, grammar, formatting and to preserve technical keywords.
Keep the resume concise and professional. Return ONLY the resume content (no commentary)."""

st.header("Step 3 — Enhance Resume with AI")
enhance_key = f"enhance_{st.session_state.resume_text}_{use_gemini}"

if st.button("✨ Enhance Resume"):
    if enhance_key not in st.session_state:
        with st.spinner("Enhancing resume..."):
            try:
                enhanced_text = cached_ai_enhance(st.session_state.resume_text, enhance_prompt, use_gemini)
                enhanced_text, missing_keywords = re_inject_keywords(st.session_state.resume_text, enhanced_text)
                st.session_state.enhanced_text = enhanced_text
                st.session_state.enhance_keywords_missing = missing_keywords
                # cache enhanced
                st.session_state[enhance_key] = enhanced_text

                # update scores once
                orig_score = ai_score
                final_score = cached_ai_score(enhanced_text, use_gemini)
                st.session_state.score_history.append({"orig": orig_score, "final": final_score})

                st.success(f"Enhanced! New ATS Score: {final_score}")
                if missing_keywords:
                    st.info(f"Re-injected keywords: {', '.join(missing_keywords)}")
            except Exception as e:
                st.error(f"AI enhancement failed: {e}")
    else:
        st.session_state.enhanced_text = st.session_state[enhance_key]

if st.session_state.enhanced_text:
    st.subheader("Enhanced Resume (Preview)")
    st.text_area("Enhanced Resume", value=st.session_state.enhanced_text, height=320)

# -------------------------
# Step 4 — Template Selection
# -------------------------
st.header("Step 4 — Template Selection")
template_choice = st.selectbox("Choose a LaTeX template:", list(TEMPLATE_MAP.keys()))
selected_template_file = TEMPLATE_MAP[template_choice]
st.write(f"Selected template file: `{selected_template_file}`")

# -------------------------
# Step 5 — Generate & Download
# -------------------------
st.header("Step 5 — Generate Resume (PDF / DOCX)")
final_text = st.session_state.enhanced_text or st.session_state.resume_text
orig_live = cached_ai_score(st.session_state.resume_text, use_gemini)
final_score = cached_ai_score(final_text, use_gemini)
st.metric("Original ATS Score", f"{orig_live} / 100", delta=None)
st.metric("Final ATS Score", f"{final_score} / 100", delta=(final_score - orig_live))

col_download_pdf, col_download_docx = st.columns(2)
with col_download_pdf:
    if st.button("📄 Generate PDF from LaTeX"):
        with st.spinner("Filling template and compiling LaTeX..."):
            fields = extract_sections(final_text)
            pdf_path = build_latex_resume(selected_template_file, fields)
            if pdf_path:
                with open(pdf_path, "rb") as f: pdf_bytes = f.read()
                st.success("PDF generated using LaTeX!")
                st.download_button("📥 Download LaTeX PDF", data=pdf_bytes, file_name="AI_Enhanced_Resume.pdf", mime="application/pdf")
            else:
                pdf_buf = generate_pdf_from_text(final_text)
                st.warning("LaTeX compile failed or not available — offering fallback PDF.")
                st.download_button("📥 Download Fallback PDF", data=pdf_buf, file_name="AI_Enhanced_Resume.pdf", mime="application/pdf")

with col_download_docx:
    if st.button("📥 Download DOCX (plain)"):
        docx_bytes = generate_docx_from_text(final_text)
        st.download_button("Download DOCX", data=docx_bytes, file_name="AI_Enhanced_Resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# -------------------------
# Comparison & History
# -------------------------
st.header("Comparison & Score Tracker")
c1, c2 = st.columns(2)
with c1: st.text_area("Original Resume", value=st.session_state.resume_text, height=300)
with c2: st.text_area("Enhanced Resume", value=st.session_state.enhanced_text or st.session_state.resume_text, height=300)

if st.session_state.score_history:
    st.subheader("Score Improvement History")
    rows = st.session_state.score_history[-12:]
    st.line_chart({"original": [r["orig"] for r in rows], "final": [r["final"] for r in rows]})

# -------------------------
# Sidebar Feedback Chat
# -------------------------
st.sidebar.header("💬 Quick Feedback Chat")
if st.session_state.resume_text:
    msg = st.sidebar.text_area("Ask for feedback:", height=90)
    if st.sidebar.button("Send feedback request"):
        with st.sidebar.spinner("Getting feedback..."):
            try:
                if use_gemini and GEMINI_KEY:
                    model = genai.GenerativeModel("models/gemini-2.5-flash")
                    resp = model.generate_content(f"Provide concise actionable feedback on this resume:\n\n{st.session_state.enhanced_text or st.session_state.resume_text}")
                    st.sidebar.write(resp.text)
                elif OPENAI_KEY:
                    client = OpenAI(api_key=OPENAI_KEY)
                    r = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role":"user","content":f"Provide concise actionable feedback on this resume:\n\n{st.session_state.enhanced_text or st.session_state.resume_text}"}],
                        temperature=0.3,
                        max_tokens=450
                    )
                    st.sidebar.write(r.choices[0].message.content)
                else:
                    st.sidebar.info("No AI key configured.")
            except Exception as e:
                st.sidebar.error(f"Feedback call failed: {e}")








# -------------------------
# Sidebar Feedback Chat
# -------------------------
st.sidebar.header("💬 Quick Feedback Chat")
if st.session_state.resume_text:
    msg = st.sidebar.text_area("Ask for feedback:", height=90)
    if st.sidebar.button("Send feedback request"):
        with st.sidebar.spinner("Getting feedback..."):
            try:
                if use_gemini and GEMINI_KEY:
                    model = genai.GenerativeModel("models/gemini-2.5-flash")
                    resp = model.generate_content(f"Provide concise actionable feedback on this resume:\n\n{st.session_state.enhanced_text or st.session_state.resume_text}")
                    st.sidebar.write(resp.text)
                elif OPENAI_KEY:
                    client = OpenAI(api_key=OPENAI_KEY)
                    r = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role":"user","content":f"Provide concise actionable feedback on this resume:\n\n{st.session_state.enhanced_text or st.session_state.resume_text}"}],
                        temperature=0.3,
                        max_tokens=450
                    )
                    st.sidebar.write(r.choices[0].message.content)
                else:
                    st.sidebar.info("No AI key configured.")
            except Exception as e:
                st.sidebar.error(f"Feedback call failed: {e}")




